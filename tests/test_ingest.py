"""Ingest round-trip tests on synthetic fixtures built at test time.

Fixtures are generated, so each carries the SYNTHETIC header in the body per
workspace policy.
"""

import docx
import pymupdf
import pytest

from deadline_agent.ingest import NoTextLayerError, UnsupportedFormatError, ingest
from deadline_agent.segment import segment

SYNTHETIC_HEADER = "SYNTHETIC — GENERATED TEST DATA"

CONTRACT_LINES = [
    SYNTHETIC_HEADER,
    "ARTICLE 4 — CLAIMS AND DISPUTES",
    "4.1 Notice of Claims. Claims by either party must be initiated by",
    "written notice to the other party within 21 days after occurrence of",
    "the event giving rise to such Claim.",
    "4.2 Differing Site Conditions. The Contractor shall give written",
    "notice within 7 days of discovery of any concealed or unknown",
    "condition, and shall submit a priced claim within 21 days thereafter.",
]


@pytest.fixture
def sample_pdf(tmp_path):
    path = tmp_path / "sample.pdf"
    doc = pymupdf.open()
    page = doc.new_page()
    y = 72
    for line in CONTRACT_LINES:
        page.insert_text((72, y), line, fontsize=11)
        y += 18
    doc.save(str(path))
    doc.close()
    return path


@pytest.fixture
def sample_docx(tmp_path):
    path = tmp_path / "sample.docx"
    d = docx.Document()
    for line in CONTRACT_LINES:
        d.add_paragraph(line)
    d.save(str(path))
    return path


def test_pdf_roundtrip(sample_pdf):
    result = ingest(sample_pdf)
    assert result.kind == "pdf"
    assert result.page_count == 1
    assert all(b.page == 1 for b in result.blocks)
    # PyMuPDF's base font substitutes the em dash, so match the words only.
    assert "SYNTHETIC" in result.blocks[0].text
    assert "GENERATED TEST DATA" in result.blocks[0].text

    clauses = segment(result.blocks)
    by_number = {c.number: c for c in clauses}
    assert "4.1" in by_number and "4.2" in by_number
    assert by_number["4.1"].heading == "Notice of Claims"
    assert "within 21 days after occurrence" in by_number["4.1"].text
    assert "within 7 days of discovery" in by_number["4.2"].text


def test_docx_roundtrip(sample_docx):
    result = ingest(sample_docx)
    assert result.kind == "docx"
    assert result.page_count is None

    clauses = segment(result.blocks)
    by_number = {c.number: c for c in clauses}
    assert by_number["4.2"].heading == "Differing Site Conditions"
    assert "priced claim within 21 days" in by_number["4.2"].text


def test_docx_list_numbering_is_captured(tmp_path):
    path = tmp_path / "listed.docx"
    d = docx.Document()
    d.add_paragraph(SYNTHETIC_HEADER)
    d.add_paragraph("First numbered item", style="List Number")
    d.save(str(path))
    result = ingest(path)
    assert result.blocks[1].is_list_item


def test_scanned_pdf_raises(tmp_path):
    path = tmp_path / "scanned.pdf"
    doc = pymupdf.open()
    doc.new_page()  # a page with no text layer
    doc.save(str(path))
    doc.close()
    with pytest.raises(NoTextLayerError):
        ingest(path)


def test_unsupported_format(tmp_path):
    path = tmp_path / "contract.txt"
    path.write_text("not supported")
    with pytest.raises(UnsupportedFormatError):
        ingest(path)
